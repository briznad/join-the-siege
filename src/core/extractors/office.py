from typing import List, Optional
from .base import BaseExtractor, ExtractedContent
import docx
import openpyxl
from openpyxl.utils import get_column_letter
import pandas as pd
import logging
from ...exceptions.classification import ExtractionError

logger = logging.getLogger(__name__)

class WordExtractor(BaseExtractor):
    @property
    def supported_mimes(self) -> List[str]:
        return [
            'application/msword',  # .doc
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document'  # .docx
        ]

    def extract_content(self, file_path: str) -> ExtractedContent:
        try:
            doc = docx.Document(file_path)

            # Extract main text
            text = []
            for paragraph in doc.paragraphs:
                text.append(paragraph.text)

            # Extract headers and footers
            headers = []
            footers = []
            for section in doc.sections:
                headers.extend([header.text for header in section.header.paragraphs])
                footers.extend([footer.text for footer in section.footer.paragraphs])

            # Extract tables
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    table_data.append([cell.text for cell in row.cells])
                tables.append(table_data)

            # Collect metadata
            metadata = {
                'page_count': len(doc.sections),
                'paragraph_count': len(doc.paragraphs),
                'table_count': len(tables),
                'word_count': len(' '.join(text).split()),
                'has_headers': bool(headers),
                'has_footers': bool(footers)
            }

            # Create final content
            final_text = '\n'.join(text)

            return ExtractedContent(
                text=self._clean_text(final_text),
                metadata=metadata,
                tables=tables,
                headers=headers,
                footers=footers,
                language=self._detect_language(final_text),
                confidence=self._calculate_confidence(final_text)
            )

        except Exception as e:
            logger.error(f"Word extraction error: {str(e)}", exc_info=True)
            raise ExtractionError(f"Failed to extract Word content: {str(e)}")

    def validate_file(self, file_path: str) -> bool:
        try:
            docx.Document(file_path)
            return True
        except Exception:
            return False

class ExcelExtractor(BaseExtractor):
    @property
    def supported_mimes(self) -> List[str]:
        return [
            'application/vnd.ms-excel',  # .xls
            'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'  # .xlsx
        ]

    def extract_content(self, file_path: str) -> ExtractedContent:
        try:
            workbook = openpyxl.load_workbook(file_path, data_only=True)

            text_content = []
            all_tables = []
            headers = []

            for sheet in workbook.worksheets:
                # Convert sheet to DataFrame for easier processing
                data = []
                for row in sheet.iter_rows():
                    data.append([str(cell.value) if cell.value is not None else ''
                               for cell in row])
                df = pd.DataFrame(data)

                # Detect tables within the sheet
                tables = self._detect_tables(df)
                all_tables.extend(tables)

                # Extract header rows
                if len(data) > 0:
                    headers.append(data[0])

                # Extract text content
                text_content.append(f"Sheet: {sheet.title}")
                for row in data:
                    text_content.extend([str(cell) for cell in row if cell])

            metadata = {
                'sheet_count': len(workbook.worksheets),
                'table_count': len(all_tables),
                'total_rows': sum(sheet.max_row for sheet in workbook.worksheets),
                'total_columns': sum(sheet.max_column for sheet in workbook.worksheets)
            }

            final_text = '\n'.join(text_content)

            return ExtractedContent(
                text=self._clean_text(final_text),
                metadata=metadata,
                tables=all_tables,
                headers=headers,
                language=self._detect_language(final_text),
                confidence=self._calculate_confidence(final_text)
            )

        except Exception as e:
            logger.error(f"Excel extraction error: {str(e)}", exc_info=True)
            raise ExtractionError(f"Failed to extract Excel content: {str(e)}")

    def validate_file(self, file_path: str) -> bool:
        try:
            openpyxl.load_workbook(file_path, data_only=True)
            return True
        except Exception:
            return False

    def _detect_tables(self, df: pd.DataFrame) -> List[List[str]]:
        """Detect table structures within a DataFrame."""
        tables = []
        current_table = []

        for idx, row in df.iterrows():
            if not row.isna().all():  # If row is not empty
                current_table.append(row.tolist())
            elif current_table:  # Empty row after table content
                if len(current_table) > 1:  # Minimum table size
                    tables.append(current_table)
                current_table = []

        if current_table and len(current_table) > 1:
            tables.append(current_table)

        return tables

    def _is_header_row(self, row: List[str]) -> bool:
        """Determine if a row is likely a header row."""
        non_empty = [cell for cell in row if cell]
        if not non_empty:
            return False

        # Check if the row contains typical header keywords
        header_keywords = {'total', 'sum', 'average', 'qty', 'amount', 'price', 'date'}
        row_text = ' '.join(str(cell).lower() for cell in row)
        return any(keyword in row_text for keyword in header_keywords)
