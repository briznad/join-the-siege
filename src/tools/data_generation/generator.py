from typing import Dict, List, Optional
from datetime import datetime
import random
from faker import Faker
from docx import Document
from docx.shared import Inches
import os
import logging

logger = logging.getLogger(__name__)

class DocumentGenerator:
    def __init__(self, output_dir: str):
        self.faker = Faker()
        self.output_dir = output_dir
        self._setup_directories()

    def _setup_directories(self):
        """Create necessary directories for generated files."""
        for industry in ['financial', 'healthcare']:
            path = os.path.join(self.output_dir, industry)
            os.makedirs(path, exist_ok=True)

    def generate_dataset(
        self,
        num_documents: int,
        industry_distribution: Optional[Dict[str, float]] = None
    ) -> List[dict]:
        """Generate a diverse dataset of documents."""
        if not industry_distribution:
            industry_distribution = {
                'financial': 0.5,
                'healthcare': 0.5
            }

        metadata = []
        for _ in range(num_documents):
            industry = random.choices(
                list(industry_distribution.keys()),
                weights=list(industry_distribution.values()),
                k=1
            )[0]

            doc_meta = self._generate_document(industry)
            metadata.append(doc_meta)

        return metadata

    def _generate_document(self, industry: str) -> dict:
        """Generate a single document based on industry."""
        generators = {
            'financial': self._generate_financial_document,
            'healthcare': self._generate_healthcare_document
        }

        generator = generators.get(industry)
        if not generator:
            raise ValueError(f"Unknown industry: {industry}")

        return generator()

    def _generate_financial_document(self) -> dict:
        """Generate a random financial document."""
        doc_types = ['bank_statement', 'invoice']
        doc_type = random.choice(doc_types)

        if doc_type == 'bank_statement':
            return self._generate_bank_statement()
        else:
            return self._generate_invoice()

    def _generate_bank_statement(self) -> dict:
        """Generate a realistic bank statement."""
        doc = Document()

        # Add header
        doc.add_heading('Bank Statement', 0)
        doc.add_paragraph(f'Statement Period: {self.faker.date()} - {self.faker.date()}')

        # Add account information
        account_info = doc.add_paragraph()
        account_info.add_run('Account Number: ').bold = True
        account_info.add_run(f'****{self.faker.random_number(digits=4)}')

        # Generate transactions
        transactions_added = 0
        balance = random.uniform(1000, 10000)

        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        header_cells = table.rows[0].cells
        headers = ['Date', 'Description', 'Amount', 'Balance']
        for i, header in enumerate(headers):
            header_cells[i].text = header

        for _ in range(random.randint(15, 30)):
            amount = random.uniform(-500, 1000)
            balance += amount

            row_cells = table.add_row().cells
            row_cells[0].text = str(self.faker.date_this_month())
            row_cells[1].text = self._generate_transaction_description()
            row_cells[2].text = f"${amount:.2f}"
            row_cells[3].text = f"${balance:.2f}"
            transactions_added += 1

        # Save document
        filename = f"bank_statement_{self.faker.random_number(digits=6)}.docx"
        filepath = os.path.join(self.output_dir, 'financial', filename)
        doc.save(filepath)

        return {
            'industry': 'financial',
            'type': 'bank_statement',
            'filename': filename,
            'filepath': filepath,
            'metadata': {
                'transaction_count': transactions_added,
                'final_balance': balance
            }
        }

    def _generate_invoice(self) -> dict:
        doc = Document()

        # Add header
        doc.add_heading('Invoice', 0)

        # Add company information
        company_info = doc.add_paragraph()
        company_info.add_run('From:\n').bold = True
        company_info.add_run(f'{self.faker.company()}\n')
        company_info.add_run(f'{self.faker.address()}\n')

        # Add invoice details
        details = doc.add_paragraph()
        details.add_run('Invoice Number: ').bold = True
        details.add_run(f'INV-{self.faker.random_number(digits=6)}\n')
        details.add_run('Date: ').bold = True
        details.add_run(f'{self.faker.date()}\n')

        # Add items
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        header_cells = table.rows[0].cells
        for i, header in enumerate(['Item', 'Quantity', 'Price', 'Total']):
            header_cells[i].text = header

        total = 0
        items_added = 0
        for _ in range(random.randint(3, 8)):
            qty = random.randint(1, 10)
            price = random.uniform(10, 1000)
            item_total = qty * price
            total += item_total

            row_cells = table.add_row().cells
            row_cells[0].text = self.faker.word()
            row_cells[1].text = str(qty)
            row_cells[2].text = f"${price:.2f}"
            row_cells[3].text = f"${item_total:.2f}"
            items_added += 1

        # Add total
        doc.add_paragraph(f'Total: ${total:.2f}')

        # Save document
        filename = f"invoice_{self.faker.random_number(digits=6)}.docx"
        filepath = os.path.join(self.output_dir, 'financial', filename)
        doc.save(filepath)

        return {
            'industry': 'financial',
            'type': 'invoice',
            'filename': filename,
            'filepath': filepath,
            'metadata': {
                'total_amount': total,
                'item_count': items_added
            }
        }

    def _generate_healthcare_document(self) -> dict:
        """Generate a healthcare document."""
        doc_types = ['medical_record', 'prescription', 'lab_report']
        doc_type = random.choice(doc_types)

        if doc_type == 'medical_record':
            return self._generate_medical_record()
        elif doc_type == 'prescription':
            return self._generate_prescription()
        else:
            return self._generate_lab_report()

    def _generate_medical_record(self) -> dict:
        """Generate a medical record document."""
        doc = Document()

        # Add header
        doc.add_heading('Medical Record', 0)

        # Add patient information
        patient = doc.add_paragraph()
        patient.add_run('Patient Information\n').bold = True
        patient.add_run(f'Name: {self.faker.name()}\n')
        patient.add_run(f'DOB: {self.faker.date_of_birth().strftime("%Y-%m-%d")}\n')
        patient.add_run(f'MRN: {self.faker.random_number(digits=8)}\n')

        # Add vital signs
        doc.add_heading('Vital Signs', level=1)
        vitals_table = doc.add_table(rows=1, cols=2)
        vitals_table.style = 'Table Grid'
        rows_added = 0

        vitals = [
            ('Blood Pressure', f'{random.randint(110, 140)}/{random.randint(60, 90)} mmHg'),
            ('Heart Rate', f'{random.randint(60, 100)} bpm'),
            ('Temperature', f'{random.uniform(97.0, 99.0):.1f}°F'),
            ('Respiratory Rate', f'{random.randint(12, 20)} /min'),
            ('Weight', f'{random.randint(120, 200)} lbs')
        ]
        for vital, value in vitals:
            row_cells = vitals_table.add_row().cells
            row_cells[0].text = vital
            row_cells[1].text = value
            rows_added += 1

        # Save document
        filename = f"medical_record_{self.faker.random_number(digits=6)}.docx"
        filepath = os.path.join(self.output_dir, 'healthcare', filename)
        doc.save(filepath)

        return {
            'industry': 'healthcare',
            'type': 'medical_record',
            'filename': filename,
            'filepath': filepath,
            'metadata': {
                'mrn': str(self.faker.random_number(digits=8)),
                'vital_signs_count': rows_added
            }
        }

    def _generate_prescription(self) -> dict:
        """Generate a prescription document."""
        doc = Document()

        # Add header
        doc.add_heading('Prescription', 0)

        # Add prescription details
        rx = doc.add_paragraph()
        rx.add_run('Rx\n').bold = True
        rx.add_run(f'Date: {self.faker.date_this_month()}\n\n')
        rx.add_run(f'Patient: {self.faker.name()}\n')
        rx.add_run(f'DOB: {self.faker.date_of_birth().strftime("%Y-%m-%d")}\n\n')

        # Add medication
        med = doc.add_paragraph()
        med.add_run(f'{self.faker.word().capitalize()} {random.randint(5, 500)}mg\n')
        med.add_run(f'Sig: Take 1 tablet by mouth {random.choice(["daily", "twice daily", "three times daily"])}\n')
        med.add_run(f'Disp: #{random.randint(30, 90)} tablets\n')
        med.add_run(f'Refills: {random.randint(0, 3)}\n')

        # Save document
        filename = f"prescription_{self.faker.random_number(digits=6)}.docx"
        filepath = os.path.join(self.output_dir, 'healthcare', filename)
        doc.save(filepath)

        return {
            'industry': 'healthcare',
            'type': 'prescription',
            'filename': filename,
            'filepath': filepath,
            'metadata': {}
        }

    def _generate_lab_report(self) -> dict:
        """Generate a lab report document."""
        doc = Document()

        # Add header
        doc.add_heading('Laboratory Report', 0)

        # Add patient information
        patient = doc.add_paragraph()
        patient.add_run('Patient Information\n').bold = True
        patient.add_run(f'Name: {self.faker.name()}\n')
        patient.add_run(f'DOB: {self.faker.date_of_birth().strftime("%Y-%m-%d")}\n')
        patient.add_run(f'Collection Date: {self.faker.date_this_month()}\n')

        # Add results table
        doc.add_heading('Test Results', level=1)
        results_table = doc.add_table(rows=1, cols=4)
        results_table.style = 'Table Grid'

        # Add headers
        header_cells = results_table.rows[0].cells
        headers = ['Test Name', 'Result', 'Units', 'Reference Range']
        for i, header in enumerate(headers):
            header_cells[i].text = header

        # Add test results
        rows_added = 0
        tests = [
            ('Glucose', f'{random.randint(70, 120)}', 'mg/dL', '70-100'),
            ('Hemoglobin', f'{random.uniform(12, 16):.1f}', 'g/dL', '12-16'),
            ('WBC', f'{random.uniform(4, 11):.1f}', 'K/uL', '4.5-11.0'),
            ('Platelets', f'{random.randint(150, 400)}', 'K/uL', '150-400')
        ]

        for test in tests:
            row_cells = results_table.add_row().cells
            for i, value in enumerate(test):
                row_cells[i].text = str(value)
            rows_added += 1

        # Save document
        filename = f"lab_report_{self.faker.random_number(digits=6)}.docx"
        filepath = os.path.join(self.output_dir, 'healthcare', filename)
        doc.save(filepath)

        return {
            'industry': 'healthcare',
            'type': 'lab_report',
            'filename': filename,
            'filepath': filepath,
            'metadata': {
                'test_count': rows_added
            }
        }

    def _generate_transaction_description(self) -> str:
        """Generate realistic transaction descriptions."""
        templates = [
            f"POS DEBIT {self.faker.company()} {self.faker.city()}",
            f"ACH CREDIT {self.faker.company()} PAYROLL",
            f"ONLINE TRANSFER TO {self.faker.name()}",
            f"ATM WITHDRAWAL {self.faker.city()}",
            f"CHECK #{random.randint(1000, 9999)}",
            f"DEPOSIT #{random.randint(1000, 9999)}",
            f"BILL PAY TO {self.faker.company()}"
        ]
        return random.choice(templates)

    def generate_test_files(self, count: int = 3) -> Dict[str, List[str]]:
        """Generate a set of test files for each document type."""
        test_files = {
            'bank_statements': [],
            'invoices': [],
            'medical_records': [],
            'prescriptions': [],
            'lab_reports': []
        }

        for _ in range(count):
            # Generate bank statement
            bank_doc = self._generate_bank_statement()
            test_files['bank_statements'].append(bank_doc['filepath'])

            # Generate invoice
            invoice_doc = self._generate_invoice()
            test_files['invoices'].append(invoice_doc['filepath'])

            # Generate medical record
            medical_doc = self._generate_medical_record()
            test_files['medical_records'].append(medical_doc['filepath'])

            # Generate prescription
            prescription_doc = self._generate_prescription()
            test_files['prescriptions'].append(prescription_doc['filepath'])

            # Generate lab report
            lab_doc = self._generate_lab_report()
            test_files['lab_reports'].append(lab_doc['filepath'])

        return test_files